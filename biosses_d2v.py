import os
from subprocess import run


class BIOSSESDataset:
    """
    Dataset class for the BIOSSES Benchmark.
    The BIOSSES dataset includes 100 biomedical sentence pairs from
    the TAC 2014 Biomedical Summarization Track (https://tac.nist.gov/2014/BiomedSumm/)
    whose similarity was judged by 5 different human experts according to
    the SemEval 2012 Task 6 Guideline on a scale of 0-4.
    This class can prepare the BIOSSES dataset, both the sentence pairs
    and similarity scores, in array form.
    Source: https://tabilab.cmpe.boun.edu.tr/BIOSSES/DataSet.html.
    """

    import docx
    import spacy
    import scipy
    from pandas import DataFrame

    def __init__(self):
        """Downloads BIOSSES data if the current directory does not contain it."""

        if not os.path.isfile("BIOSSES-Dataset/Annotation-Pairs.docx") \
           or not os.path.isfile("BIOSSES-Dataset/Annotator-Scores.docx"):
            run(["wget", "https://tabilab.cmpe.boun.edu.tr/\
                          BIOSSES/Downloads/BIOSSES-Dataset.rar"])
            run(["unrar", "x", "BIOSSES-Dataset.rar"])

    def _get_df_from_doc(self, path):
        """Loads a table from a .doc/docx file into an array.

        Args: 
          path: Filepath of the document.

        Returns:
          DataFrame obtained from the file. 
        """

        doc = self.docx.Document(path)
        table = doc.tables[0]
        df = self.DataFrame([[cell.text.strip() for cell in row.cells]
                             for row in table.rows])
        return df

    def get_sentence_df(self):
        """Gets the BIOSSES sentence pairs for benchmarking."""

        sent_df = self._get_df_from_doc(
            "BIOSSES-Dataset/Annotation-Pairs.docx")
        return sent_df.iloc[1:, 1:].rename(columns={1: "Sentence_1",
                                                    2: "Sentence_2"})

    def get_score_df(self):
        """Gets the BIOSSES similarity scores by expert annotators."""

        score_df = self._get_df_from_doc(
            "BIOSSES-Dataset/Annotator-Scores.docx")
        return score_df.iloc[1:, 1:].astype("i1")

    def benchmark_with_d2v(self, d2v_model=None, lemma=False, stopwords=[]):
        """Benchmarks a Doc2Vec model with the BIOSSES Dataset. 
        Returns the Pearson correlation with annotators' similarity scores 
        and also the cosine similarity scores produced by the Doc2Vec model.
        """

        cosine = self.scipy.spatial.distance.cosine
        pearsonr = self.scipy.stats.pearsonr
        nlp = self.spacy.load("en_core_sci_sm", disable=["parser", "ner"])

        def preprocess(sentence):
            if lemma:
                tokens = [token.lemma_ for token in nlp(sentence)
                          if token not in stopwords]
            else:
                tokens = [token for token in sentence.split()
                          if token not in stopwords]
            return tokens

        def similarity(sent1, sent2):
            tokens1 = preprocess(sent1)
            tokens2 = preprocess(sent2)
            d2v_model.random.seed(0)
            return 1 - cosine(d2v_model.infer_vector(tokens1),
                              d2v_model.infer_vector(tokens2))

        bio_sents = self.get_sentence_df()
        real_scores = self.get_score_df()
        d2v_scores = bio_sents.apply(
            lambda row: similarity(row[0], row[1]), axis=1)
        avg_real_scores = real_scores.apply(
            lambda row: sum(row)/len(row), axis=1)
        corr, _ = pearsonr(avg_real_scores, d2v_scores)

        return corr, d2v_scores


class PMCOASubsetCorpus:
    """
    Corpus iterator for the PMC Open Access (PMCOA) Subset in plain text form.
    PMCOA Subset is a corpus of select PubMed Central articles made free for 
    research purposes and is accessible via FTP server (ftp://ftp.ncbi.nlm.nih.gov/pub/pmc).
    The iterator generates such text data through the oa_bulk directory, from
    both usage groups: commercial and non-commercial use.
    Source: https://www.ncbi.nlm.nih.gov/pmc/tools/ftp/.
    """

    import re
    import spacy
    import smart_open
    from gensim.models.doc2vec import TaggedDocument

    def __init__(self,
                 packages=["0-9A-B"],
                 size=float("inf"),
                 lemma=False,
                 stopwords=[],
                 sub_pattern=None):
        """Loads user-defined corpus packages.
        User can determine the number of articles from all or any combination 
        of the following packages: 0-9A-B, C-H, I-N, O-Z.
        Preprocessing is optional in the following ways:
          - lemmatization
          - stopword removal
          - regex substitution
        """

        self.packages = packages
        self.size = size
        self.lemma = lemma
        self.stopwords = stopwords
        self.sub_pattern = sub_pattern
        # Could try r"(\s+\(*(([À-ÿA-Za-z\s\-.,;&])+\s\(*(\d{4}[a-z]*)+\)*)+)|(\s+\[[\d\s+,;&\[\]]+\])"
        # for getting rid of majority of in-text citations
        self.corpus_paths = self._load()
        self.nlp = self.spacy.load(
            name="en_core_sci_sm", disable=["parser", "ner"])

    def __iter__(self):
        """Enables users to treat this object as an iterable."""

        return self._iterator()

    def _iterator(self):
        """Creates a generator of preprocessed articles from the PMCOA Subset."""

        doc_id = 0
        for paper_path in self.corpus_paths:
            tagged_doc = self._preprocess_doc(paper_path, doc_id)
            if tagged_doc is not None:
                yield tagged_doc
                doc_id += 1
                if doc_id == self.size:
                    break

    def _preprocess_doc(self, path, id):
        """Tokenizes text from a file.

        Args:
          path: Filepath of the article to be preprocessed.
          id: Article ID number required for training Doc2Vec.

        Returns:
          Tokens in the form of a TaggedDocument,
          following Doc2Vec's training protocol.
        """

        with self.smart_open.open(path, "r", encoding="ascii", errors="ignore") as f:
            paper = f.read()
            if self.sub_pattern is not None:
                paper = self.re.sub(self.sub_pattern, "", paper)
            body = self.re.search(
                "==== Body(.*)==== Refs", paper, self.re.DOTALL)

            if body is not None:
                doc = body.group(1).lower().split()
                if self.lemma:
                    doc = " ".join(doc)
                    doc = [token.lemma_ for token in self.nlp(doc)
                           if not token in self.stopwords]
                else:
                    doc = [token for token in doc
                           if not token in self.stopwords]
                return self.TaggedDocument(doc, [id])

    def _append_corpus_paths(self, start_path, path_arr):
        """Recursively scans and appends paths of all .txt files 
        from a directory and its subdirectories, if any, to a provided array. 

        Args:
          start_path: Directory path to be scanned for .txt files.
          path_arr: Array to append all the .txt paths.
        """

        for entry in os.scandir(start_path):
            if not entry.name.startswith(".") \
               and entry.is_dir(follow_symlinks=False):
                self._append_corpus_paths(entry.path, path_arr)

            if entry.name.endswith(".txt"):
                path_arr.append(entry.path)

    def _load(self):
        """Downloads the zipped corpus folders and unzips
        them if they are not already available in the current directory.
        Then loads individual filepaths of all .txt files into 
        an internal array, necessary for the iterator later.
        """

        corpus_paths = []

        for package in self.packages:
            zipnames = ["non_comm_use." + package + ".txt.tar.gz",
                        "comm_use." + package + ".txt.tar.gz"]

            for zname in zipnames:
                dirname = zname.replace(".txt.tar.gz", "")
                if not os.path.isdir(dirname):
                    os.mkdir(dirname)
                    if not os.path.isfile(zname):
                        run(["wget",
                            "ftp://ftp.ncbi.nlm.nih.gov/pub/pmc/oa_bulk/"
                             + zname])
                    run(["tar", "-xf", zname, "-C", dirname])
                self._append_corpus_paths(dirname, corpus_paths)

        return corpus_paths

    def get_list(self):
        doc_list = []
        doc_id = 0
        for paper_path in self.corpus_paths:
            tagged_doc = self._preprocess_doc(paper_path, doc_id)
            if tagged_doc is not None:
                doc_list.append(tagged_doc)
                doc_id += 1
                if doc_id == self.size:
                    break

        return doc_list


class Doc2VecRunner:
    """
    Runner to train Doc2Vec from the Gensim library on a corpus.
    This is just an abstraction to streamline Doc2Vec's configurations.
    It enables training logs and model saving if need be. 
    Doc2Vec's documentation: https://radimrehurek.com/gensim/models/doc2vec.html.
    """

    import logging
    from gensim.models.doc2vec import Doc2Vec

    def __init__(self, corpus=None, **kwargs):
        """Initializes internal model and corpus"""

        self.model = self.Doc2Vec(**kwargs)
        self.corpus = corpus

    def run(self,
            use_logger=False,
            log_dir=None,
            progress_per=None):
        """Trains the Doc2Vec model, with optional progress logging"""

        if use_logger:
            log_dir = "model" if log_dir is None else log_dir
            log_fname = "{}.log".format(log_dir)
            if not os.path.isdir(log_dir):
                os.mkdir(log_dir)
            logpath = os.path.join(log_dir, log_fname)
            self.logging.basicConfig(
                format="%(asctime)s : %(levelname)s : %(message)s",
                level=self.logging.INFO,
                handlers=[
                    self.logging.FileHandler(filename=logpath,
                                             mode="w"),
                    self.logging.StreamHandler()
                ])

        if progress_per is not None:
            self.model.build_vocab(self.corpus,
                                   progress_per=progress_per)
        else:
            self.model.build_vocab(self.corpus)
        self.model.train(self.corpus,
                         total_examples=self.model.corpus_count,
                         epochs=self.model.epochs)
        return self.model

    def save_model(self, model_dir=None):
        """Saves model's internal weights to a directory"""

        model_dir = "model" if model_dir is None else model_dir
        model_fname = "{}.gsm".format(model_dir)
        if not os.path.isdir(model_dir):
            os.mkdir(model_dir)
        savepath = os.path.join(model_dir, model_fname)
        self.model.save(savepath)


def run_doc2vec(iterator,
                lemma,
                use_logger,
                save_model,
                package_names,
                corpus_size,
                progress_per,
                **kwargs):
    """Program to run based on CLI arguments.
    Helps train a configurable Doc2Vec model on a configurable corpus.
    If the model is saved, the default saving directory is determined 
    based on whether lemmatization or data iteration is enabled.
    """

    corpus = PMCOASubsetCorpus(
        packages=package_names, size=corpus_size, lemma=lemma)
    if not iterator:
        corpus = corpus.get_list()

    model_dir = "biod2v_{}lemma_{}iter".format(
        "!" if not lemma else "", "!" if not iterator else "")
    runner = Doc2VecRunner(corpus, **kwargs)
    model = runner.run(use_logger, model_dir, progress_per=progress_per)
    if save_model:
        runner.save_model(model_dir)
    biosses = BIOSSESDataset()
    correlation, _ = biosses.benchmark_with_d2v(model, lemma)
    print("Correlation:", correlation)


def main():
    """
    Parsing arguments to power Doc2Vec with the PMCOA Subset corpus
    if user wants to use this module in the CLI.
    """

    import argparse
    from multiprocessing import cpu_count

    parser = argparse.ArgumentParser()
    parser.add_argument(
        "-i", "--iterator",
        action="store_true",
        help="Whether data is streamed in using an iterator or \
              stored in memory in a list object."
    )
    parser.add_argument(
        "-l", "--lemma",
        action="store_true",
        help="Whether tokens are turned into lemmas for training."
    )
    parser.add_argument(
        "--use-logger",
        action="store_true",
        help="Whether to show training progress."
    )
    parser.add_argument(
        "--save-model",
        action="store_true",
        help="Whether to save the model weights."
    )
    parser.add_argument(
        "-n", "--package-names",
        type=str,
        default=["0-9A-B"],
        help="Names of PMC Open Access packages used for training."
    )
    parser.add_argument(
        "-s", "--corpus-size",
        type=float,
        default=float("inf"),
        help="Size (number of articles) of the training corpus."
    )
    parser.add_argument(
        "-p", "--progress-per",
        type=int,
        default=1000,
        help="Number of words processed before showing progress."
    )
    parser.add_argument(
        "-v", "--vector-size",
        type=int,
        default=100,
        help="Dimensionality of Doc2Vec embeddings."
    )
    parser.add_argument(
        "-w", "--window",
        type=int,
        default=5,
        help="Maximum distance between the current and predicted word \
              during training."
    )
    parser.add_argument(
        "-a", "--alpha",
        type=float,
        default=0.025,
        help="Initial learning rate."
    )
    parser.add_argument(
        "-m", "--min-count",
        type=int,
        default=5,
        help="Minimum frequency of accepted words."
    )
    parser.add_argument(
        "-W", "--workers",
        type=int,
        default=cpu_count(),
        help="Number of worker threads used to train the model."
    )
    parser.add_argument(
        "-e", "--epochs",
        type=int,
        default=10,
        help="Number of iterations trained over the corpus."
    )
    args, _ = parser.parse_known_args()
    run_doc2vec(iterator=args.iterator,
                lemma=args.lemma,
                use_logger=args.use_logger,
                save_model=args.save_model,
                package_names=args.package_names,
                corpus_size=args.corpus_size,
                progress_per=args.progress_per,
                vector_size=args.vector_size,
                window=args.window,
                alpha=args.alpha,
                min_count=args.min_count,
                workers=args.workers,
                epochs=args.epochs)


if __name__ == "__main__":
    main()
