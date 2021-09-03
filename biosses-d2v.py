import os
from subprocess import run
from numpy import array
from docx import Document

class BIOSSESDataset:
    """
    Dataset class for the BIOSSES Benchmark.
    Source: https://tabilab.cmpe.boun.edu.tr/BIOSSES/DataSet.html.
    """
    
    def __init__(self):
        BIOSSES_URL = "https://tabilab.cmpe.boun.edu.tr/BIOSSES/Downloads/BIOSSES-Dataset.rar"
        if not os.path.isfile("BIOSSES-Dataset/Annotation-Pairs.docx") \
          or not os.path.isfile("BIOSSES-Dataset/Annotator-Scores.docx"):
            run(["wget", BIOSSES_URL])
            run(["unrar", "x", "BIOSSES-Dataset.rar"])

 
    def get_sentences(self):
        doc = Document(os.path.abspath("BIOSSES-Dataset/Annotation-Pairs.docx"))
        table = doc.tables[0]
        sent_arr = array([[cell.text.strip() for cell in row.cells] for row in table.rows], dtype="i1")
        return sent_arr

    
    def get_scores(self):
        doc = Document(os.path.abspath("BIOSSES-Dataset/Annotator-Scores.docx"))
        table = doc.tables[0]
        score_arr = array([[cell.text.strip() for cell in row.cells] for row in table.rows])
        return score_arr


class PMCSubsetCorpus:
    def __init__(self):
        self.corpus_paths = []
        self.cite_pattern = r'(\s+\(*(([À-ÿA-Za-z\s\-.,;&])+\s\(*(\d{4}[a-z]*)+\)*)+)|(\s+\[[\d\s+,;&\[\]]+\])'
        self.doc_id = 0


    def __iter__(self):
        for paper_path in self.corpus_paths:
            with open(paper_path, 'r', encoding='ascii', errors='ignore') as f:
                paper = re.sub(self.cite_pattern, '', f.read())
                body = re.search('==== Body(.*)==== Refs', paper, re.DOTALL)

                if body is not None:
                    body = body.group(1).split()
                    doc = [token for token in body if not token in STOPWORDS]
                    yield TaggedDocument(doc, [self.doc_id])
                    self.doc_id += 1


    def _append_corpus_paths(self, start_path, path_arr):
        for entry in os.scandir(start_path):
            if not entry.name.startswith('.') \
              and not entry.name == 'drive' \
              and not entry.name == 'sample_data' \
              and entry.is_dir(follow_symlinks=False):
                  _append_corpus_paths(entry.path, path_arr)

            if not entry.name.startswith('.') and entry.is_file():
                path_arr.append(entry.path)


    def download(self, subsets=["0-9A-B"]):
        PMCSC_URL = "ftp://ftp.ncbi.nlm.nih.gov/pub/pmc/oa_bulk/"
        ZIP_NAMES = []

        for subset in subsets:
            ZIP_NAMES.append("non_comm_use." + subset + ".txt.tar.gz")

        for zname in ZIP_NAMES:
            fdir = zname.replace(".txt.tar.gz", "")
            fdir_path = os.path.abspath(fdir)
            if not os.path.isdir(fdir_path):
                run(["mkdir", fdir])

                if not os.path.isfile(os.path.abspath(zname)):
                    run(["wget", PMCSC_URL + zname])

                run(["tar", "-xf", zname, "-C", fdir])
                _append_corpus_paths(fdir_path, self.corpus_paths)

